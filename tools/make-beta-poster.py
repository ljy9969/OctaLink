"""
체육관 벽보용 A4 포스터 (.docx) 생성 — OctaLink 비공개 테스터 모집.

산출물: docs/beta-tester-poster.docx
QR 이미지는 Word 에서 "삽입 > 그림" 으로 가운데 [QR 코드 영역] 자리에 직접 추가.
선택적으로 --qr 옵션에 PNG 경로 주면 자동 삽입.

사용:
    python tools/make-beta-poster.py
    python tools/make-beta-poster.py --qr docs/openchat-qr.png
"""

import argparse
import os
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


FONT_NAME = "Malgun Gothic"  # 한글 가독성 + Windows 기본 폰트
BRAND_RED = RGBColor(0xC8, 0x10, 0x2E)
SLATE = RGBColor(0x1A, 0x1A, 0x1F)
PEBBLE = RGBColor(0x6A, 0x6A, 0x72)


def set_font(run, name=FONT_NAME, size=None, bold=False, color=None):
    """한글 + 영문 모두 같은 폰트 적용. python-docx 가 기본으로 east-asian 폰트 별도 설정 필요."""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text, *, size=11, bold=False, color=None, align=None, space_before=0, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def add_divider(doc):
    """수평선 — paragraph 의 아래 border 로 구현."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C8102E")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_poster(output_path: str, qr_path: str | None = None):
    doc = Document()

    # A4 portrait, 여백 1.8cm
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    # ─────────────────────────────────────
    # 헤더 — 타이틀 + 부제
    # ─────────────────────────────────────
    add_paragraph(
        doc,
        "OctaLink",
        size=44,
        bold=True,
        color=BRAND_RED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "Team Posse Striking 강남점 회원 전용 앱",
        size=14,
        color=SLATE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "비공개 테스터 모집 (선착순 12명)",
        size=18,
        bold=True,
        color=SLATE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6,
        space_after=6,
    )

    add_divider(doc)

    # ─────────────────────────────────────
    # 왜 테스터가 필요한가
    # ─────────────────────────────────────
    add_paragraph(
        doc,
        "왜 테스터가 필요한가요?",
        size=14,
        bold=True,
        color=BRAND_RED,
        space_before=8,
        space_after=4,
    )
    intro = (
        "Google Play Store 공식 출시를 위해서는 비공개 테스트를 거쳐야 합니다. "
        "체육관 회원 12명이 함께 사용해보고 의견을 주시면, 정식 출시 후 모든 회원이 "
        "더 안정적인 앱을 받아볼 수 있습니다."
    )
    add_paragraph(doc, intro, size=11, color=SLATE, space_after=10)

    # ─────────────────────────────────────
    # 주요 기능
    # ─────────────────────────────────────
    add_paragraph(
        doc,
        "주요 기능",
        size=14,
        bold=True,
        color=BRAND_RED,
        space_before=4,
        space_after=4,
    )
    features = [
        ("카카오 로그인 + 자동 가입", "별도 비밀번호 없이 간편 시작"),
        ("셀프 출석 체크 + 동료 출석", "수업 30분 전후 자동 활성, 함께한 회원 한눈에"),
        ("토너먼트 추첨 / 대진표 / 전적", "체급·벨트 필터 추첨, 챔피언 폭죽 연출, 히스토리 누적"),
        ("6축 스킬 차트 + 관장 한 줄 코멘트", "스트라이킹·그래플링·체력·기술·멘탈·스피드 시각화"),
        ("커뮤니티 (글 / 이미지 / 영상)", "기록·팁·질문 공유, 30초 영상 첨부, 좋아요"),
        ("주간 미션 + 푸시 알림", "수업 리마인더, 코멘트·공지·대진표 도착 알림"),
        ("다크 / 라이트 테마 선택", "기기별 선호 저장"),
    ]
    for title, desc in features:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        # 큰 점 + 제목
        run_dot = p.add_run("● ")
        set_font(run_dot, size=11, bold=True, color=BRAND_RED)
        run_title = p.add_run(title)
        set_font(run_title, size=11, bold=True, color=SLATE)
        run_desc = p.add_run(f"  —  {desc}")
        set_font(run_desc, size=10, color=PEBBLE)

    # ─────────────────────────────────────
    # 참여 방법 — QR
    # ─────────────────────────────────────
    add_paragraph(
        doc,
        "참여 방법",
        size=14,
        bold=True,
        color=BRAND_RED,
        space_before=14,
        space_after=4,
    )
    add_paragraph(
        doc,
        "아래 QR 코드를 카카오톡으로 스캔하면 오픈채팅방에 입장합니다.",
        size=11,
        color=SLATE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "채팅방 운영자가 테스트 설치 링크 / 가이드를 안내합니다.",
        size=10,
        color=PEBBLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    # QR 영역 — 표 1셀로 placeholder 박스 만들거나 이미지 삽입
    qr_paragraph = doc.add_paragraph()
    qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if qr_path and os.path.isfile(qr_path):
        qr_paragraph.add_run().add_picture(qr_path, width=Cm(7.5))
    else:
        # 자리 표시 — 점선 박스를 시각화하기 어려우니 명확한 안내 텍스트.
        run = qr_paragraph.add_run("[ 여기에 QR 이미지 삽입 — 권장 7~8cm ]")
        set_font(run, size=12, color=PEBBLE)
        qr_paragraph.paragraph_format.space_before = Pt(20)
        qr_paragraph.paragraph_format.space_after = Pt(20)

    add_paragraph(
        doc,
        "▼ 카카오 오픈채팅 입장 QR",
        size=10,
        color=PEBBLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=2,
        space_after=14,
    )

    # ─────────────────────────────────────
    # 안내사항
    # ─────────────────────────────────────
    add_paragraph(
        doc,
        "안내사항",
        size=12,
        bold=True,
        color=SLATE,
        space_before=4,
        space_after=4,
    )
    notes = [
        "현재 안드로이드 OS 만 지원합니다. (iOS 는 추후 검토)",
        "테스트 중 발견한 오류 / 불편 / 개선 아이디어는 오픈채팅으로 알려주세요.",
        "선착순 12명 마감 후에도 정식 출시까지 1~2주 정도 소요됩니다.",
        "테스트 참여자 개인정보는 별도 수집하지 않으며, 카카오 로그인 정보만 사용합니다.",
    ]
    for note in notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        run_b = p.add_run("- ")
        set_font(run_b, size=10, color=PEBBLE)
        run_n = p.add_run(note)
        set_font(run_n, size=10, color=SLATE)

    # ─────────────────────────────────────
    # 푸터
    # ─────────────────────────────────────
    add_divider(doc)
    add_paragraph(
        doc,
        "OctaLink · Unbound Apex Systems · 개발 BlackCat Strike",
        size=9,
        color=PEBBLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=4,
        space_after=0,
    )

    doc.save(output_path)
    print(f"saved: {output_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--qr", help="QR 이미지 PNG 경로 (선택, 없으면 placeholder)")
    parser.add_argument(
        "--out",
        default=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs", "beta-tester-poster.docx"),
        help="출력 .docx 경로",
    )
    args = parser.parse_args()
    build_poster(args.out, args.qr)
